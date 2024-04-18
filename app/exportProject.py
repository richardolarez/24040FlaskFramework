from app import db
from models import Projects, ChargeMode, Devices, GSENetwork, PathsLoads, PowerBusConfig, PowerSupply, PowerSupplyAssign, PowerSupplySummary, TelemetryNetwork, VehicleBattery, VehicleNetwork, UEIDaq, BatteryAddresses, BatteryDefault, TIDTables, Component, ExternalMode
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm

def export_Project(projectId, file_path):
    project = Projects.query.filter_by(id=projectId).first()
    doc = Document()
    PowerSupplyTables(doc)
    PathsAndLoadsTable(doc)
    ChargeModeTables(doc)
    GseNetIpAddresses(doc)
    BatteryAddressTable(doc)
    SamplingRateTable(doc)
    VehicleNet(doc)
    BatteryDefaultTable(doc)
    BatteryDischargeTable(doc)
    PowerSupplyAssignTable(doc)
    ExternalModeTables(doc)
    PowerBusConfigTable(doc)
    daqs = Component.query.filter_by(componentType = '[DAQ-Digital]').all()
    for daq in daqs: 
         UEIDaqTable(doc, daq)
    doc.save(file_path)

def PowerSupplyTables(doc):
    power_supplies = PowerSupply.query.all()
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Power Supply"
    hdr_cells[1].text = "Battery/System"
    hdr_cells[2].text = "Voltage Setting (V)"
    hdr_cells[3].text = "OVP (V)"
    hdr_cells[4].text = "Current Limit (A)"
    hdr_cells[-2].merge(hdr_cells[-1])  # Merge the last two cells for the first header
    hdr_cells[-2].text = 'Red/Green Limits'
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
        # Add the additional header row
    additional_header_row = table.add_row().cells
    additional_header_row[5].text = 'Voltage'
    additional_header_row[6].text = 'Current'
    for i in range(5):
        table.cell(0, i).merge(table.cell(1, i))
    # Add data rows to the table
    for power_supply in power_supplies:
        row_cells = table.add_row().cells
        row_cells[0].text = power_supply.power_supply
        row_cells[1].text = power_supply.battery_system
        row_cells[2].text = str(power_supply.voltage_setting)
        row_cells[3].text = str(power_supply.ovp)
        row_cells[4].text = str(power_supply.current_limit)
        row_cells[5].text = power_supply.red_green_voltage_limits
        row_cells[6].text = power_supply.red_green_current_limits
    doc.add_paragraph() 

def ChargeModeTables(doc):
    chargeTable = ChargeMode.query.all()
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Power Supply"
    hdr_cells[1].text = "Battery/System"
    hdr_cells[2].text = "Voltage Setting (V)"
    hdr_cells[3].text = "OVP (V)"
    hdr_cells[4].text = "Current Settinng (A)"
    hdr_cells[5].text = "Current Limit (A)"
    hdr_cells[-2].merge(hdr_cells[-1])  # Merge the last two cells for the first header
    hdr_cells[-2].text = 'Red/Green Limits'
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
        # Add the additional header row
    additional_header_row = table.add_row().cells
    additional_header_row[6].text = 'Voltage'
    additional_header_row[7].text = 'Current'
    for i in range(6):
        table.cell(0, i).merge(table.cell(1, i))
    # Add data rows to the table
    for charge in chargeTable:
        row_cells = table.add_row().cells
        row_cells[0].text = charge.power_supply
        row_cells[1].text = charge.battery
        row_cells[2].text = str(charge.voltage_setting)
        row_cells[3].text = str(charge.ovp)
        row_cells[4].text = str(charge.current_setting)
        row_cells[5].text = str(charge.current_limit)
        row_cells[6].text = charge.red_green_voltage_limits
        row_cells[7].text = charge.red_green_current_limits
    doc.add_paragraph() 

def GseNetIpAddresses(doc):
    networks = GSENetwork.query.all()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "GSE NET DEVICE"
    hdr_cells[1].text = "IP ADDRESS"
    hdr_cells[2].text = "SUB NET MASK"
    hdr_cells[3].text = "HOST NAME"

    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
    for network in networks:
        row_cells = table.add_row().cells
        row_cells[0].text = network.gse_net_device
        row_cells[1].text = network.ip_address
        row_cells[2].text = network.sub_net_mask
        row_cells[3].text = network.host_name

    doc.add_paragraph() 

def ChargeModeTables(doc):
    chargeTable = ChargeMode.query.all()
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Power Supply"
    hdr_cells[1].text = "Battery/System"
    hdr_cells[2].text = "Voltage Setting (V)"
    hdr_cells[3].text = "OVP (V)"
    hdr_cells[4].text = "Current Settinng (A)"
    hdr_cells[5].text = "Current Limit (A)"
    hdr_cells[-2].merge(hdr_cells[-1])  # Merge the last two cells for the first header
    hdr_cells[-2].text = 'Red/Green Limits'
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
        # Add the additional header row
    additional_header_row = table.add_row().cells
    additional_header_row[6].text = 'Voltage'
    additional_header_row[7].text = 'Current'
    for i in range(6):
        table.cell(0, i).merge(table.cell(1, i))
    # Add data rows to the table
    for charge in chargeTable:
        row_cells = table.add_row().cells
        row_cells[0].text = charge.power_supply
        row_cells[1].text = charge.battery
        row_cells[2].text = str(charge.voltage_setting)
        row_cells[3].text = str(charge.ovp)
        row_cells[4].text = str(charge.current_setting)
        row_cells[5].text = str(charge.current_limit)
        row_cells[6].text = charge.red_green_voltage_limits
        row_cells[7].text = charge.red_green_current_limits
    doc.add_paragraph() 

def ExternalModeTables(doc):
    chargeTable = ExternalMode.query.all()
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Power Supply"
    hdr_cells[1].text = "Battery/System"
    hdr_cells[2].text = "Voltage Setting (V)"
    hdr_cells[3].text = "OVP (V)"
    hdr_cells[4].text = "Current Settinng (A)"
    hdr_cells[5].text = "Current Limit (A)"
    hdr_cells[-2].merge(hdr_cells[-1])  # Merge the last two cells for the first header
    hdr_cells[-2].text = 'Red/Green Limits'
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
        # Add the additional header row
    additional_header_row = table.add_row().cells
    additional_header_row[6].text = 'Voltage'
    additional_header_row[7].text = 'Current'
    for i in range(6):
        table.cell(0, i).merge(table.cell(1, i))
    # Add data rows to the table
    for charge in chargeTable:
        row_cells = table.add_row().cells
        row_cells[0].text = charge.power_supply
        row_cells[1].text = charge.battery
        row_cells[2].text = str(charge.voltage_setting)
        row_cells[3].text = str(charge.ovp)
        row_cells[4].text = str(charge.current_setting)
        row_cells[5].text = str(charge.current_limit)
        row_cells[6].text = charge.red_green_voltage_limits
        row_cells[7].text = charge.red_green_current_limits
    doc.add_paragraph() 


def BatteryAddressTable(doc):
    batteries = BatteryAddresses.query.all()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "BATTERY"
    hdr_cells[1].text = "Serial Address"

    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
    for battery in batteries:
        row_cells = table.add_row().cells
        row_cells[0].text = battery.battery
        row_cells[1].text = battery.rs485_address

    doc.add_paragraph() 
     
def SamplingRateTable(doc):
    devices = Devices.query.all()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "DEVICE"
    hdr_cells[1].text = "Sampling Rate"

    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.device
        row_cells[1].text = device.sampling_rate

    doc.add_paragraph() 

def VehicleNet(doc):
    devices = VehicleNetwork.query.all()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "DEVICE"
    hdr_cells[1].text = "IP Address"

    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.device
        row_cells[1].text = device.ip_address
        
    doc.add_paragraph()

def BatteryDefaultTable(doc):
    devices = BatteryDefault.query.all()
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "BATTERY"
    hdr_cells[1].text = "CAPACITY"
    hdr_cells[2].text = "DISCHARGE CURRENT (A)"

    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.battery
        row_cells[1].text = device.capacity
        row_cells[2].text = device.discharge_current
        
    doc.add_paragraph()
     
def BatteryDischargeTable(doc):
    devices = VehicleBattery.query.all()
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "BATTERY"
    hdr_cells[1].text = "PS V (V)"
    hdr_cells[2].text = "UEI V (V)"
    hdr_cells[3].text = "BAT V (V)"
    hdr_cells[4].text = "CELL (V)"
    hdr_cells[5].text = "TEMP (DEG C)"
    hdr_cells[6].text = "LOAD V (V)"
    hdr_cells[7].text = "LOAD I (A)"
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.battery
        row_cells[1].text = device.psv
        row_cells[2].text = device.ueiv
        row_cells[3].text = device.batv
        row_cells[4].text = device.cell
        row_cells[5].text = device.temp
        row_cells[6].text = device.loadv
        row_cells[7].text = device.loadi

    doc.add_paragraph() 

def PowerSupplyAssignTable(doc):
    devices = PowerSupplyAssign.query.all()
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "POWER SUPPLY"
    hdr_cells[1].text = "BATTERY/SYSTEM"
    hdr_cells[2].text = "DEVICES"
    hdr_cells[3].text = "EXT PWR"
    hdr_cells[4].text = "BATT CHG"
    hdr_cells[5].text = "CONTROL"
    hdr_cells[6].text = "MONITOR"
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.power_supply
        row_cells[1].text = device.battery
        row_cells[2].text = device.devices
        row_cells[3].text = device.ext_pwr
        row_cells[4].text = device.batt_chg
        row_cells[5].text = device.control
        row_cells[6].text = device.monitor

    doc.add_paragraph() 

def PowerBusConfigTable(doc):
    devices = PowerBusConfig.query.all()
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "PWR SPLY"
    hdr_cells[1].text = "BATTERY"
    hdr_cells[2].text = "COMPONENT"
    hdr_cells[3].text = "EXT PWR"
    hdr_cells[4].text = "INT PWR"
    hdr_cells[-4].merge(hdr_cells[-1])  # Merge the last 4 cells for the first header
    hdr_cells[-4].text = 'Red/Green Limits'
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
        # Add the additional header row
    additional_header_row = table.add_row().cells
    additional_header_row[5].text = 'BUS V LOW'
    additional_header_row[6].text = 'BUS V HIGH'
    additional_header_row[7].text = 'BUS I LOW'
    additional_header_row[8].text = 'BUS I HIGH'

    for i in range(5):
        table.cell(0, i).merge(table.cell(1, i))
    # Add data rows to the table
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.power_supply
        row_cells[1].text = device.battery
        row_cells[2].text = device.component
        row_cells[3].text = device.ext_pwr
        row_cells[4].text = device.int_pwr
        row_cells[5].text = device.bus_v_low
        row_cells[6].text = device.bus_v_high
        row_cells[7].text = device.bus_i_low
        row_cells[8].text = device.bus_i_high
    doc.add_paragraph() 

def UEIDaqTable(doc, daq):
    devices = UEIDaq.query.filter_by(power_daq_layer=daq.componentName).all()
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    title_cells = table.rows[0].cells 
    title_cells[-4].merge(title_cells[-1])
    title_cells[-4].text = str(daq.componentName)
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = "BIT"
    hdr_cells[1].text = "PIN"
    hdr_cells[2].text = "SIGNAL"
    hdr_cells[3].text = "INITIAL"
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.bit
        row_cells[1].text = device.pin
        row_cells[2].text = device.signal
        row_cells[3].text = device.initial
    doc.add_paragraph()

def PathsAndLoadsTable(doc):
    devices = PathsLoads.query.all()
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "PWR SPLY"
    hdr_cells[1].text = "BATTERY"
    hdr_cells[2].text = "PCD Channel"
    hdr_cells[3].text = "Components"
    hdr_cells[4].text = "Current"
    hdr_cells[5].text = 'Range'
   
    # Add borders to all cells in the table
    for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'

    # Add data rows to the table
    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = device.power_supply
        row_cells[1].text = device.battery
        row_cells[2].text = device.ptm_channel
        row_cells[3].text = device.components
        row_cells[4].text = str(device.current)
        row_cells[5].text = device.range_
    doc.add_paragraph() 