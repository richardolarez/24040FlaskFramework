# -*- coding: utf-8 -*-
"""
Created on Tue Feb 13 16:35:51 2024

@author: julia
"""

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm

class TableGenerator: 
    def __init__(self, document):
        self.document = document
        
    def generate_table(self, num_rows, data):
        # Add a table with 7 columns and a header row
        table = self.document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'  # Add grid style to the table
    
        # Apply bold style to the header row
        hdr_cells = table.rows[0].cells
    
        hdr_cells[0].text = "Power Supply"
        hdr_cells[1].text = "Battery/System"
        hdr_cells[2].text = "Voltage Setting (V)"
        hdr_cells[3].text = "OVP (V)"
        hdr_cells[4].text = "Current Limit (A)"
        hdr_cells[-2].merge(hdr_cells[-1])  # Merge the last two cells for the first header
        hdr_cells[-2].text = 'Red/Green Limits'
   
        # Add borders to all cells in the table
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
    
        # Add the additional header row
        additional_header_row = table.add_row().cells
        additional_header_row[5].text = 'Voltage'
        additional_header_row[6].text = 'Current'
        for i in range(5):
            table.cell(0, i).merge(table.cell(1, i))
            # Add the specified number of rows
        for _ in range(num_rows):
            row_cells = table.add_row().cells
            row_cells[0].text = data[_]
        
        self.document.add_paragraph()
        
    def generate_small_table(self, num_rows, data):
        # Add a table with 4 columns and a header row
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'  # Add grid style to the table
        
        # Apply bold style to the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "GSE Net Device"
        hdr_cells[1].text = "IP Address"
        hdr_cells[2].text = "Sub Net Mask"
        hdr_cells[3].text = "Host Name"
        
        
        # Add borders to all cells in the table
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                for side in cell._element.xpath('.//w:tcBorders/w:*'):
                    side.attrib.clear()
                    side.attrib['w:val'] = 'single'
                    side.attrib['w:sz'] = '4'
                    side.attrib['w:color'] = 'auto'
        
        # Add the specified number of rows
        for _ in range(num_rows):
            row_cells = table.add_row().cells
            row_cells[0].text = data[_]
         
        self.document.add_paragraph()


doc = Document()
ps = ["POWER SUPPLY 1", "POWER SUPPLY 2", "POWER SUPPLY 3"]
data2 = ["UEI DAQ 1", "UEI DAQ 2", "CSE Server ETH0", "PDU", "POWER SUPPLY 1", "POWER SUPPLY 2", "POWER SUPPLY 3"]
table_generator = TableGenerator(doc)
table_generator.generate_table(len(ps), ps)
table_generator.generate_small_table(len(data2), data2)
doc.save('table_document.docx')
print("Table document generated successfully.")